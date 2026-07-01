from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

NAVY        = RGBColor(0x1B, 0x2A, 0x4A)
DARK_GRAY   = RGBColor(0x2D, 0x2D, 0x2D)
MID_GRAY    = RGBColor(0x55, 0x55, 0x55)
ACCENT_HEX  = "1B2A4A"
BODY_FONT   = "Calibri"
NAME_SIZE   = 22
SECTION_SIZE= 11
BODY_SIZE   = 10.5
CONTACT_SIZE= 10


def _make_border_element(tag, val="single", sz="12", color=ACCENT_HEX, space="4"):
    el = OxmlElement(tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    return el


def add_accent_underline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pBdr.append(_make_border_element("w:bottom", sz="18", color=ACCENT_HEX, space="2"))
    pPr.append(pBdr)


def add_section_header(doc, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pBdr.append(_make_border_element("w:left", sz="24", color=ACCENT_HEX, space="6"))
    pBdr.append(_make_border_element("w:bottom", sz="4", color="CCCCCC", space="2"))
    pPr.append(pBdr)

    run = p.add_run(title.upper())
    run.font.name  = BODY_FONT
    run.font.size  = Pt(SECTION_SIZE)
    run.font.bold  = True
    run.font.color.rgb = NAVY


def add_company_line(doc, company: str, role: str, dates: str):
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)

    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),   "right")
    tab.set(qn("w:pos"),   "9360")
    tabs.append(tab)
    pPr.append(tabs)

    r_company = p.add_run(company)
    r_company.font.name  = BODY_FONT
    r_company.font.size  = Pt(BODY_SIZE)
    r_company.font.bold  = True
    r_company.font.color.rgb = DARK_GRAY

    r_tab = p.add_run("\t")
    r_tab.font.size = Pt(BODY_SIZE)

    r_date = p.add_run(dates)
    r_date.font.name   = BODY_FONT
    r_date.font.size   = Pt(BODY_SIZE - 0.5)
    r_date.font.italic = True
    r_date.font.color.rgb = MID_GRAY

    if role:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(2)
        r_role = p2.add_run(role)
        r_role.font.name   = BODY_FONT
        r_role.font.size   = Pt(BODY_SIZE)
        r_role.font.italic = True
        r_role.font.color.rgb = MID_GRAY


def add_bullet(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.2)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    clean = text.lstrip("-•* ").strip()
    run = p.add_run(f"- {clean}")
    run.font.name  = BODY_FONT
    run.font.size  = Pt(BODY_SIZE)
    run.font.color.rgb = DARK_GRAY


def add_body(doc, text: str, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name   = BODY_FONT
    run.font.size   = Pt(BODY_SIZE)
    run.font.italic = italic
    run.font.color.rgb = DARK_GRAY


KNOWN_HEADERS = {
    "summary", "objective", "profile", "about",
    "skills", "technical skills", "core competencies",
    "experience", "work experience", "professional experience",
    "education",
    "projects", "personal projects", "key projects",
    "certifications", "certificates", "awards"
}

def parse_cv(cv_text: str) -> dict:
    lines    = cv_text.strip().split("\n")
    sections = {}
    current  = "HEADER"
    buf      = []

    for line in lines:
        s     = line.strip()
        lower = s.lower().rstrip(":")
        if lower in KNOWN_HEADERS:
            sections[current] = buf
            current = lower.upper().rstrip(":")
            buf = []
        else:
            buf.append(s)

    sections[current] = buf
    return sections


def fmt_contact(lines: list) -> str:
    parts = []
    for item in lines:
        s   = item.strip(" |,")
        if not s:
            continue
        low = s.lower()
        if "@" in s:
            parts.append(f"✉ {s}")
        elif "linkedin" in low:
            parts.append(f"\U0001f4bc {s}")
        elif "github" in low:
            parts.append(f"\U0001f517 {s}")
        elif re.match(r"[\+\d][\d\s\-\(\)]{6,}", s):
            parts.append(f"\U0001f4de {s}")
        elif any(c in low for c in ["cairo", "egypt", "new york", "london", "remote", "city"]):
            parts.append(f"⌂ {s}")
        else:
            parts.append(s)
    return "   |   ".join(parts)


def parse_experience_blocks(lines: list) -> list:
    blocks  = []
    current = None

    for line in lines:
        if not line.strip():
            continue
        if re.search(r"\b(19|20)\d{2}\b", line) and not line.startswith("-"):
            if current:
                blocks.append(current)
            current = {"raw": line, "bullets": []}
        elif line.startswith("-") or line.startswith("•"):
            if current:
                current["bullets"].append(line)
            else:
                current = {"raw": "", "bullets": [line]}
        else:
            if current and not current["bullets"]:
                current["raw"] += " " + line
            elif current:
                current["bullets"].append(line)
            else:
                current = {"raw": line, "bullets": []}

    if current:
        blocks.append(current)
    return blocks


def split_company_line(raw: str):
    date_match = re.search(
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|\d{4})[\w\s,\-–]+(?:Present|\d{4}))",
        raw, re.IGNORECASE
    )
    dates     = date_match.group(1).strip() if date_match else ""
    remainder = raw.replace(dates, "").strip(" |-,")
    parts     = re.split(r"\s*[-|,]\s*", remainder, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip(), dates
    return remainder.strip(), "", dates


SECTION_ORDER = [
    "SUMMARY", "OBJECTIVE", "PROFILE", "ABOUT",
    "EDUCATION",
    "SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES",
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
    "PROJECTS", "PERSONAL PROJECTS", "KEY PROJECTS",
    "CERTIFICATIONS", "CERTIFICATES", "AWARDS"
]


def generate_cv_docx(cv_text: str, output_path: str, candidate_name: str = "") -> str:
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)

    sections = parse_cv(cv_text)

    header_lines = [l for l in sections.get("HEADER", []) if l.strip()]
    name = candidate_name or (header_lines[0] if header_lines else "")

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(name)
        run.font.name  = BODY_FONT
        run.font.size  = Pt(NAME_SIZE)
        run.font.bold  = True
        run.font.color.rgb = NAVY

    add_accent_underline(doc)

    contact_items = header_lines[1:] if header_lines else []
    contact_str   = fmt_contact(contact_items)
    if contact_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(contact_str)
        run.font.name  = BODY_FONT
        run.font.size  = Pt(CONTACT_SIZE)
        run.font.color.rgb = MID_GRAY

    rendered = set()

    for key in SECTION_ORDER:
        if key in sections and key not in rendered:
            lines = [l for l in sections[key] if l.strip()]
            if not lines:
                continue
            rendered.add(key)

            title = key.title().replace("And", "and")
            add_section_header(doc, title)

            if key in {"EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE"}:
                blocks = parse_experience_blocks(lines)
                for block in blocks:
                    if block["raw"].strip():
                        company, role, dates = split_company_line(block["raw"])
                        add_company_line(doc, company, role, dates)
                    for bullet in block["bullets"]:
                        add_bullet(doc, bullet)
            else:
                for line in lines:
                    if line.startswith("-") or line.startswith("•") or line.startswith("*"):
                        add_bullet(doc, line)
                    else:
                        add_body(doc, line)

    dir_name = os.path.dirname(output_path)
    if dir_name:
        os.makedirs(dir_name, exist_ok=True)
    doc.save(output_path)
    return output_path
